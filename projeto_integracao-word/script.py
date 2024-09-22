from docx import Document
from docx.shared import Cm, RGBColor, Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime



documento = Document()

for estilo in documento.styles:
    print(estilo)

for secao in documento.sections:
    secao.top_margin = Cm(0.5)
    secao.right_margin = Cm(1)
    secao.left_margin = Cm(1)
    secao.bottom_margin = Cm(1)


proventos = 104.56

t1 = f"""Olá Lucas, Tudo bem?

Os proventos desse mês foram de {proventos:,.2f}.

Continue firme.

Abraços!
"""
p1 = documento.add_paragraph(t1)

p1.style = documento.styles.add_style("Estilo Lucas", WD_STYLE_TYPE.PARAGRAPH)
p1.style.font.name = "Algerian"
p1.style.font.size = Pt(15)
p1.style.font.bold = False
p1.style.font.italic = True
p1.style.font.underline = False

investido = 10000
p2 = documento.add_paragraph(f"Vale ressalta que você possui investido o total de ", "Estilo Lucas")
p2.add_run(f"R${investido:,.2f}").bold = True

p3 = documento.add_picture(r"Python/projeto_integracao-word/imagem.png", width=Cm(5), height=Cm(5))
p3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

records = {
    (1, "100", "Carne"),
    (5, "200", "Frango"),
    (7, "563", "Ovo")
}

table = documento.add_table(rows=1, cols=3, style="Light Grid Accent 1")
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "Qty"
hdr_cells[1].text = "ID"
hdr_cells[2].text = "Desc"
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

documento.save("teste01.docx")


contrato = Document(r"Python\projeto_integracao-word\Contrato.docx")

valor_contrato = 2000

nome = "Lucas"
item1 = "Serviço de desenvolvimento Web"
item2 = "Construção de um site com HTML, CSS e JS hospedado na Hostnet"
item3 = f"Valor Total de R${valor_contrato:,.2f}"

dicionario_valores = {
    "XXXX": nome,
    "YYYY": item1,
    "ZZZZ": item2,
    "WWWW": item3,
    "DD": str(datetime.now().day),
    "MM": str(datetime.now().month),
    "AAAA": str(datetime.now().year)
}

for paragrafo in contrato.paragraphs:
    for placeholder in dicionario_valores:
        if placeholder in paragrafo.text:
            paragrafo.text = paragrafo.text.replace(placeholder, dicionario_valores[placeholder])

contrato.save("contrato_atualizado.docx")
